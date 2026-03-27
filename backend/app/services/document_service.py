"""Document processing service: extraction, chunking, image extraction."""
import os
import io
import re
import uuid
import hashlib
import shutil
import subprocess
import tempfile
import zipfile
from typing import List, Dict, Tuple, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

from ..config import settings


# Chunk size configuration (aligned with multilingual-e5-base 512 token limit)
# ~200 words ≈ 260-300 tokens in French, well within the 512 max_seq_length.
# Smaller chunks yield more precise retrieval results.
CHUNK_SIZE = 200  # words
CHUNK_OVERLAP = 40  # words overlap between chunks

# ── Sentence-aware splitting helpers ──

# Regex to split text into sentences (French + English punctuation)
_SENTENCE_RE = re.compile(
    r'(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-Þ0-9])'  # sentence boundary
    r'|(?<=\n)\s*(?=\n)'                    # double newline (paragraph break)
)

# Regex to detect bullet-point / list-item lines
_BULLET_RE = re.compile(
    r'^\s*(?:[-–—•●◦▪▸►✓✔☑☐★⬥]'     # common bullet chars
    r'|\d{1,3}[.)]\s'                      # numbered list: "1. " or "1) "
    r'|[a-zA-Z][.)]\s'                     # lettered list: "a. " or "a) "
    r'|\[\s*[xX ]?\]\s'                    # checkbox: "[ ] " or "[x] "
    r')',
    re.MULTILINE,
)

# Regex to detect lines containing key numerical KPIs (percentages, amounts, etc.)
_KPI_RE = re.compile(
    r'\d+[.,]\d+\s*%'                      # 4,09% or 3.5%
    r'|\d+\s*%'                             # 87%
    r'|\d[\d\s]*(?:€|euros?|k€|M€|Md€)'   # monetary amounts
    r'|\d[\d\s,.]*(?:heures?|jours?|mois|ans?)\b'  # durations
    r'|\d[\d\s,.]*(?:ETP|salariés?|collaborateurs?)\b',  # headcounts
    re.IGNORECASE,
)


def _split_into_semantic_units(text: str) -> List[str]:
    """Split text into semantic units that should not be broken apart.

    A semantic unit is:
    - A paragraph (text between double newlines)
    - Within a paragraph: a group of consecutive bullet points
    - Within a paragraph: a sentence or group of short sentences

    This ensures that bullet-point lists with KPIs stay together in the same
    chunk, preventing loss of context like "4.09% de la masse salariale".
    """
    if not text.strip():
        return []

    units = []

    # First, split into paragraphs (double newline or more)
    paragraphs = re.split(r'\n\s*\n', text)

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        lines = para.split('\n')

        # Check if this paragraph is a bullet list
        bullet_lines = [l for l in lines if _BULLET_RE.match(l.strip())]
        is_bullet_list = len(bullet_lines) > len(lines) * 0.5

        if is_bullet_list:
            # Keep bullet lists as atomic units — group consecutive bullets
            # together with any preceding header line
            current_group = []
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    if current_group:
                        units.append('\n'.join(current_group))
                        current_group = []
                    continue
                current_group.append(line_stripped)

            if current_group:
                units.append('\n'.join(current_group))
        else:
            # For regular text, split into sentences but keep KPI-containing
            # sentences grouped with their surrounding context
            sentences = _SENTENCE_RE.split(para)
            sentences = [s.strip() for s in sentences if s.strip()]

            if not sentences:
                continue

            # Group sentences: if a sentence contains a KPI, keep it with
            # the previous sentence (which likely provides context)
            i = 0
            while i < len(sentences):
                group = [sentences[i]]
                # Look ahead: if next sentence has a KPI, attach it
                while (i + 1 < len(sentences)
                       and _KPI_RE.search(sentences[i + 1])
                       and len(' '.join(group + [sentences[i + 1]]).split()) < CHUNK_SIZE):
                    i += 1
                    group.append(sentences[i])
                units.append(' '.join(group))
                i += 1

    return units


def _assemble_chunks_from_units(
    units: List[str],
    document_id: str,
    document_name: str,
    category: str,
    page_number: int,
    section_title: str,
    chunk_index_start: int,
) -> List[Dict]:
    """Assemble semantic units into chunks respecting CHUNK_SIZE with overlap.

    Units are never split — if a single unit exceeds CHUNK_SIZE, it becomes
    its own chunk (better to have one oversized chunk than to break a bullet
    list with KPIs).

    Overlap is achieved by repeating the last unit(s) of a chunk at the start
    of the next chunk, so numerical data near chunk boundaries appears in both.
    """
    chunks = []
    current_units: List[str] = []
    current_word_count = 0
    idx = chunk_index_start

    for unit in units:
        unit_words = len(unit.split())

        # If adding this unit would exceed the limit, finalize current chunk
        if current_units and (current_word_count + unit_words) > CHUNK_SIZE:
            chunk_text = '\n'.join(current_units)
            if len(chunk_text.split()) >= 15:  # min viable chunk
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": chunk_text,
                    "document_id": document_id,
                    "document_name": document_name,
                    "category": category,
                    "page_number": page_number,
                    "section_title": section_title,
                    "chunk_index": idx,
                })
                idx += 1

            # Overlap: keep the last unit(s) totaling ~CHUNK_OVERLAP words
            overlap_units: List[str] = []
            overlap_words = 0
            for u in reversed(current_units):
                u_len = len(u.split())
                if overlap_words + u_len > CHUNK_OVERLAP:
                    break
                overlap_units.insert(0, u)
                overlap_words += u_len

            current_units = overlap_units
            current_word_count = overlap_words

        current_units.append(unit)
        current_word_count += unit_words

    # Flush remaining units
    if current_units:
        chunk_text = '\n'.join(current_units)
        if len(chunk_text.split()) >= 15:
            chunks.append({
                "id": str(uuid.uuid4()),
                "content": chunk_text,
                "document_id": document_id,
                "document_name": document_name,
                "category": category,
                "page_number": page_number,
                "section_title": section_title,
                "chunk_index": idx,
            })

    return chunks


class DocumentProcessor:
    """Process documents: extract text, images, and create chunks."""

    @staticmethod
    def detect_file_type(filename: str) -> str:
        """Detect file type from extension."""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        type_map = {
            "pdf": "pdf",
            "docx": "docx",
            "doc": "doc",
            "xlsx": "xlsx",
            "xls": "xls",
            "pptx": "pptx",
        }
        return type_map.get(ext, "other")

    @staticmethod
    def convert_doc_to_docx(file_content: bytes) -> bytes:
        """Convert old .doc format to .docx using LibreOffice.

        Returns the .docx file content as bytes.
        Raises RuntimeError if conversion fails.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = os.path.join(tmpdir, "input.doc")
            with open(doc_path, "wb") as f:
                f.write(file_content)

            result = subprocess.run(
                [
                    "libreoffice", "--headless", "--norestore",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    doc_path,
                ],
                capture_output=True,
                timeout=120,
            )

            docx_path = os.path.join(tmpdir, "input.docx")
            if result.returncode != 0 or not os.path.exists(docx_path):
                stderr = result.stderr.decode(errors="ignore")
                raise RuntimeError(f"LibreOffice conversion failed: {stderr[:200]}")

            with open(docx_path, "rb") as f:
                return f.read()

    @staticmethod
    def extract_text_from_doc_fallback(file_content: bytes) -> str:
        """Fallback text extraction for .doc using antiword.

        Used when LibreOffice conversion fails.
        Returns extracted text or empty string.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = os.path.join(tmpdir, "input.doc")
            with open(doc_path, "wb") as f:
                f.write(file_content)

            try:
                result = subprocess.run(
                    ["antiword", doc_path],
                    capture_output=True,
                    timeout=60,
                )
                if result.returncode == 0:
                    return result.stdout.decode("utf-8", errors="ignore")
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            # Last resort: try catdoc
            try:
                result = subprocess.run(
                    ["catdoc", doc_path],
                    capture_output=True,
                    timeout=60,
                )
                if result.returncode == 0:
                    return result.stdout.decode("utf-8", errors="ignore")
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

        return ""

    @staticmethod
    def _validate_docx(file_content: bytes) -> bool:
        """Check if file_content is a valid DOCX (ZIP with Word content type)."""
        if not zipfile.is_zipfile(io.BytesIO(file_content)):
            return False
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                if "[Content_Types].xml" not in zf.namelist():
                    return False
                ct = zf.read("[Content_Types].xml").decode("utf-8", errors="ignore")
                return "wordprocessingml" in ct
        except Exception:
            return False

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Tuple[str, int, List[Dict]]:
        """Extract text from PDF with page-level metadata.

        Returns:
            Tuple of (full_text, page_count, pages_data)
            where pages_data is list of {page_number, text, sections}
        """
        doc = fitz.open(stream=file_content, filetype="pdf")
        pages_data = []
        full_text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            full_text_parts.append(text)

            # Try to detect section headers (lines that look like titles)
            sections = []
            for line in text.split("\n"):
                stripped = line.strip()
                if (
                    stripped
                    and len(stripped) < 200
                    and not stripped.endswith(".")
                    and (stripped[0].isupper() or stripped[0].isdigit())
                    and len(stripped.split()) < 20
                ):
                    sections.append(stripped)

            pages_data.append({
                "page_number": page_num + 1,
                "text": text,
                "sections": sections,
            })

        doc.close()
        return "\n\n".join(full_text_parts), len(pages_data), pages_data

    @staticmethod
    def extract_images_from_pdf(
        file_content: bytes, document_id: str
    ) -> List[Dict]:
        """Extract images from PDF document.

        Returns list of image metadata dicts with saved file paths.
        """
        doc = fitz.open(stream=file_content, filetype="pdf")
        images = []
        images_dir = os.path.join(settings.images_dir, document_id)
        os.makedirs(images_dir, exist_ok=True)

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue

                    image_bytes = base_image["image"]
                    image_ext = base_image.get("ext", "png")
                    width = base_image.get("width", 0)
                    height = base_image.get("height", 0)

                    # Skip very small images (likely decorative)
                    if width < 50 or height < 50:
                        continue

                    # Generate filename with content hash
                    img_hash_full = hashlib.md5(image_bytes).hexdigest()
                    img_hash = img_hash_full[:8]
                    filename = f"page{page_num + 1}_img{img_index + 1}_{img_hash}.{image_ext}"
                    filepath = os.path.join(images_dir, filename)

                    # Save image
                    with open(filepath, "wb") as f:
                        f.write(image_bytes)

                    # Get surrounding text as context
                    page_text = page.get_text("text")
                    context_lines = page_text.split("\n")
                    context = " ".join(line.strip() for line in context_lines[:5] if line.strip())

                    images.append({
                        "stored_filename": filename,
                        "file_path": filepath,
                        "page_number": page_num + 1,
                        "width": width,
                        "height": height,
                        "context": context[:500],
                        "content_hash": img_hash_full,
                        "description": "",
                        "tags": [],
                    })

                except Exception as e:
                    print(f"Error extracting image from page {page_num + 1}: {e}")
                    continue

        doc.close()
        return images

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Tuple[str, List[Dict]]:
        """Extract text from DOCX with structural metadata.

        Returns:
            Tuple of (full_text, sections_data)
        """
        if not DocumentProcessor._validate_docx(file_content):
            raise ValueError(
                "Le fichier n'est pas un document Word (.docx) valide. "
                "Les fichiers .doc (ancien format) ne sont pas supportés."
            )
        doc = DocxDocument(io.BytesIO(file_content))
        full_text_parts = []
        sections_data = []
        current_section = {"title": "", "content_parts": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if para.style and para.style.name and "heading" in para.style.name.lower():
                # Save previous section
                if current_section["content_parts"]:
                    sections_data.append({
                        "title": current_section["title"],
                        "content": "\n".join(current_section["content_parts"]),
                    })
                current_section = {"title": text, "content_parts": []}
            else:
                current_section["content_parts"].append(text)

            full_text_parts.append(text)

        # Save last section
        if current_section["content_parts"]:
            sections_data.append({
                "title": current_section["title"],
                "content": "\n".join(current_section["content_parts"]),
            })

        return "\n\n".join(full_text_parts), sections_data

    @staticmethod
    def extract_images_from_docx(
        file_content: bytes, document_id: str
    ) -> List[Dict]:
        """Extract images from DOCX document."""
        if not DocumentProcessor._validate_docx(file_content):
            return []
        doc = DocxDocument(io.BytesIO(file_content))
        images = []
        images_dir = os.path.join(settings.images_dir, document_id)
        os.makedirs(images_dir, exist_ok=True)

        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    ext = content_type.split("/")[-1] if "/" in content_type else "png"
                    if ext == "jpeg":
                        ext = "jpg"

                    img_hash_full = hashlib.md5(image_data).hexdigest()
                    img_hash = img_hash_full[:8]
                    filename = f"docx_img{i + 1}_{img_hash}.{ext}"
                    filepath = os.path.join(images_dir, filename)

                    with open(filepath, "wb") as f:
                        f.write(image_data)

                    # Get image dimensions
                    try:
                        img = Image.open(io.BytesIO(image_data))
                        width, height = img.size
                    except Exception:
                        width, height = 0, 0

                    if width < 50 or height < 50:
                        continue

                    images.append({
                        "stored_filename": filename,
                        "file_path": filepath,
                        "page_number": 0,
                        "width": width,
                        "height": height,
                        "context": "",
                        "content_hash": img_hash_full,
                        "description": "",
                        "tags": [],
                    })

                except Exception as e:
                    print(f"Error extracting image from DOCX: {e}")
                    continue

        return images

    @staticmethod
    def _convert_xls_bytes_to_xlsx(file_content: bytes) -> bytes:
        """Convert old .xls format bytes to .xlsx format bytes using xlrd."""
        import xlrd
        from openpyxl import Workbook

        xls_book = xlrd.open_workbook(file_contents=file_content)
        wb = Workbook()
        wb.remove(wb.active)

        for sheet_index in range(xls_book.nsheets):
            xls_sheet = xls_book.sheet_by_index(sheet_index)
            ws = wb.create_sheet(title=xls_sheet.name)
            for row_idx in range(xls_sheet.nrows):
                for col_idx in range(xls_sheet.ncols):
                    cell_value = xls_sheet.cell_value(row_idx, col_idx)
                    cell_type = xls_sheet.cell_type(row_idx, col_idx)
                    if cell_type == xlrd.XL_CELL_DATE:
                        try:
                            date_tuple = xlrd.xldate_as_tuple(cell_value, xls_book.datemode)
                            from datetime import datetime
                            cell_value = datetime(*date_tuple)
                        except Exception:
                            pass
                    elif cell_type == xlrd.XL_CELL_BOOLEAN:
                        cell_value = bool(cell_value)
                    elif cell_type == xlrd.XL_CELL_EMPTY:
                        continue
                    ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value)

        output = io.BytesIO()
        wb.save(output)
        wb.close()
        output.seek(0)
        return output.read()

    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> tuple:
        """Extract text from Excel file, returning full text and per-sheet pages_data.

        Returns:
            Tuple of (full_text, pages_data) where pages_data treats each sheet
            as a page with the sheet name as section_title.
        """
        # Detect old .xls format by checking file signature (OLE2 magic bytes)
        # .xlsx files start with PK (ZIP), .xls files start with OLE2 compound doc header
        if file_content[:4] == b'\xd0\xcf\x11\xe0':
            file_content = DocumentProcessor._convert_xls_bytes_to_xlsx(file_content)
        wb = load_workbook(io.BytesIO(file_content), data_only=True)
        text_parts = []
        pages_data = []

        for sheet_index, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            sheet_lines = []

            for row in ws.iter_rows(values_only=True):
                row_texts = []
                for cell in row:
                    if cell is not None:
                        row_texts.append(str(cell))
                if row_texts:
                    sheet_lines.append(" | ".join(row_texts))

            sheet_text = "\n".join(sheet_lines)
            if sheet_text.strip():
                text_parts.append(f"\n=== Feuille: {sheet_name} ===\n")
                text_parts.append(sheet_text)
                pages_data.append({
                    "page_number": sheet_index,
                    "text": sheet_text,
                    "sections": [f"Feuille: {sheet_name}"],
                })

        full_text = "\n".join(text_parts)
        return full_text, pages_data

    @staticmethod
    def create_chunks(
        text: str,
        document_id: str,
        document_name: str,
        category: str,
        pages_data: Optional[List[Dict]] = None,
    ) -> List[Dict]:
        """Split text into overlapping chunks with metadata.

        Uses sentence-aware splitting to avoid breaking bullet-point lists
        or KPI data mid-item. Ensures that numerical data like "4,09% de la
        masse salariale" stays in the same chunk as its context.

        Args:
            text: Full document text
            document_id: Document UUID
            document_name: Original filename
            category: Document category
            pages_data: Optional page-level data for better chunking

        Returns:
            List of chunk dicts ready for indexing
        """
        chunks = []

        if pages_data:
            # Page-aware chunking with semantic units
            for page in pages_data:
                page_text = page.get("text", "")
                if not page_text.strip():
                    continue

                page_num = page.get("page_number", 0)
                sections = page.get("sections", [])
                current_section = sections[0] if sections else ""

                # Split into semantic units (sentences, bullet groups, KPI groups)
                units = _split_into_semantic_units(page_text)
                if not units:
                    continue

                chunks.extend(_assemble_chunks_from_units(
                    units, document_id, document_name, category,
                    page_num, current_section, len(chunks),
                ))
        else:
            # Fallback: semantic chunking without page data
            units = _split_into_semantic_units(text)
            if units:
                chunks.extend(_assemble_chunks_from_units(
                    units, document_id, document_name, category,
                    0, "", 0,
                ))

        return chunks

    @staticmethod
    def save_uploaded_file(file_content: bytes, project_id: str, filename: str) -> str:
        """Save an uploaded file to disk.

        Returns the file path.
        """
        # Sanitize project_id and filename to prevent path traversal
        safe_project_id = os.path.basename(project_id)
        safe_filename = os.path.basename(filename)
        # Remove any remaining path separators or null bytes
        safe_filename = safe_filename.replace("\x00", "").replace("/", "").replace("\\", "")
        if not safe_filename:
            safe_filename = "upload"

        project_dir = os.path.join(settings.upload_dir, safe_project_id)
        os.makedirs(project_dir, exist_ok=True)

        stored_name = f"{uuid.uuid4().hex}_{safe_filename}"
        filepath = os.path.join(project_dir, stored_name)

        # Final safety check: ensure path stays within upload_dir
        real_upload = os.path.realpath(settings.upload_dir)
        real_filepath = os.path.realpath(filepath)
        if not real_filepath.startswith(real_upload):
            raise ValueError("Path traversal detected")

        with open(filepath, "wb") as f:
            f.write(file_content)

        return filepath
