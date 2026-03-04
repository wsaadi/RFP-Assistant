"""Service for generating professional Word documents for RFP responses."""
import io
import logging
import os
import re
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image as PILImage

from ..config import settings

logger = logging.getLogger(__name__)

# Regex to match [INSERT_IMAGE:id] or [INSERT_IMAGE:id:layout] markers
_IMAGE_MARKER_RE = re.compile(r'^\s*\[INSERT_IMAGE:([^\]:]+)(?::([^\]]+))?\]\s*$')

# Valid layout modes for image insertion
_VALID_LAYOUTS = {'center', 'wrap-right', 'wrap-left', 'full-width', 'inline'}


class RFPWordService:
    """Service for generating professional RFP response Word documents."""

    @staticmethod
    def create_styles(doc: Document):
        """Create professional custom styles."""
        styles = doc.styles

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.name = "Calibri"
        h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(8)
        h1.paragraph_format.keep_with_next = True

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.name = "Calibri"
        h2.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(6)

        # Heading 3
        h3 = styles["Heading 3"]
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.name = "Calibri"
        h3.font.color.rgb = RGBColor(0x3D, 0x7A, 0xB5)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(4)

        # Normal text
        normal = styles["Normal"]
        normal.font.size = Pt(11)
        normal.font.name = "Calibri"
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        # List styles
        try:
            list_bullet = styles["List Bullet"]
            list_bullet.font.size = Pt(11)
            list_bullet.font.name = "Calibri"
            list_bullet.paragraph_format.space_after = Pt(3)
            list_bullet.paragraph_format.space_before = Pt(1)
        except KeyError:
            pass

        try:
            list_number = styles["List Number"]
            list_number.font.size = Pt(11)
            list_number.font.name = "Calibri"
            list_number.paragraph_format.space_after = Pt(3)
            list_number.paragraph_format.space_before = Pt(1)
        except KeyError:
            pass

    @staticmethod
    def add_cover_page(doc: Document, project_name: str, client_name: str,
                       rfp_reference: str, company_name: str = ""):
        """Add a professional cover page."""
        # Spacing at top
        for _ in range(4):
            doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RÉPONSE À L'APPEL D'OFFRES")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        doc.add_paragraph()

        # RFP Reference
        if rfp_reference:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Référence: {rfp_reference}")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Project name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

        for _ in range(3):
            doc.add_paragraph()

        # Separator
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

        doc.add_paragraph()

        # Client
        if client_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Client: {client_name}")
            run.font.size = Pt(14)

        # Company
        if company_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Soumissionnaire: {company_name}")
            run.font.size = Pt(14)

        for _ in range(6):
            doc.add_paragraph()

        # Confidentiality notice
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DOCUMENT CONFIDENTIEL")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

        doc.add_page_break()

    @staticmethod
    def add_table_of_contents(doc: Document):
        """Add a Table of Contents field."""
        doc.add_heading("Sommaire", level=1)

        # Add TOC field code
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u</w:instrText>')
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar)

        run = paragraph.add_run("[Mettre à jour le sommaire dans Word: clic droit > Mettre à jour les champs]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar)

        doc.add_page_break()

    @staticmethod
    def _add_inline_formatting(paragraph, text: str):
        """Parse inline markdown (bold, italic) and add formatted runs to a paragraph."""
        # Pattern to match **bold**, *italic*, ***bold italic***
        parts = re.split(r'(\*{1,3}[^*]+?\*{1,3})', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    @classmethod
    def _add_markdown_table(cls, doc: Document, lines: List[str]):
        """Parse a markdown table and add it to the document."""
        # Filter out separator lines (|---|---|)
        data_lines = []
        for line in lines:
            stripped = line.strip().strip('|')
            if stripped and not re.match(r'^[\s\-:|]+$', stripped):
                data_lines.append(line)

        if not data_lines:
            return

        # Parse cells
        rows_data = []
        for line in data_lines:
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    cls._add_inline_formatting(p, cell_text)
                    # Bold header row
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)

        # Style header row with background
        if rows_data:
            for cell in table.rows[0].cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph()  # spacing after table

    @staticmethod
    def _get_image_dimensions(filepath: str, max_width_inches: float) -> tuple:
        """Get image dimensions scaled to fit within max_width while preserving aspect ratio."""
        try:
            with PILImage.open(filepath) as img:
                w_px, h_px = img.size
            aspect = h_px / w_px if w_px > 0 else 1
            width = Inches(max_width_inches)
            height = Inches(max_width_inches * aspect)
            return width, height
        except Exception:
            return Inches(max_width_inches), None

    @classmethod
    def _insert_image_in_doc(cls, doc: Document, filepath: str, description: str = "",
                              layout: str = "center"):
        """Insert an image with professional layout into the Word document.

        Supported layouts:
        - center: Centered image (default, 4.5in wide)
        - full-width: Full page width image
        - wrap-right: Image floated right with text wrapping
        - wrap-left: Image floated left with text wrapping
        - inline: Small inline image within text flow
        """
        if not filepath or not os.path.exists(filepath):
            logger.warning("Image file not found: %s", filepath)
            return

        if layout not in _VALID_LAYOUTS:
            layout = "center"

        try:
            if layout == "full-width":
                cls._insert_image_full_width(doc, filepath, description)
            elif layout in ("wrap-right", "wrap-left"):
                cls._insert_image_wrapped(doc, filepath, description, layout)
            elif layout == "inline":
                cls._insert_image_inline(doc, filepath, description)
            else:
                cls._insert_image_centered(doc, filepath, description)
        except Exception as e:
            logger.error("Error inserting image %s (layout=%s): %s", filepath, layout, e)

    @classmethod
    def _insert_image_centered(cls, doc: Document, filepath: str, description: str):
        """Insert a centered image — classic professional layout."""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(4.5))

        if description:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Figure : {description}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    @classmethod
    def _insert_image_full_width(cls, doc: Document, filepath: str, description: str):
        """Insert a full-width image spanning the page."""
        # Calculate available width from page margins
        section = doc.sections[-1]
        avail_width = section.page_width - section.left_margin - section.right_margin
        avail_inches = avail_width / 914400  # EMU to inches

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        width, _ = cls._get_image_dimensions(filepath, avail_inches)
        run.add_picture(filepath, width=width)

        if description:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Figure : {description}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    @classmethod
    def _insert_image_wrapped(cls, doc: Document, filepath: str, description: str,
                               layout: str):
        """Insert an image with text wrapping (float left or right).

        Uses Word's wp:anchor XML to position the image with tight text wrapping.
        """
        wrap_width_inches = 2.8
        width_emu = int(wrap_width_inches * 914400)
        img_width, img_height = cls._get_image_dimensions(filepath, wrap_width_inches)
        height_emu = int(img_height) if img_height else int(width_emu * 0.75)

        # Use a table-based approach for reliable wrapping (1 row, 2 cols)
        section = doc.sections[-1]
        avail_width = section.page_width - section.left_margin - section.right_margin
        avail_inches = avail_width / 914400

        text_col_width = Inches(avail_inches - wrap_width_inches - 0.3)
        img_col_width = Inches(wrap_width_inches + 0.1)

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Remove table borders for a clean layout
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        if layout == "wrap-left":
            img_cell = table.cell(0, 0)
            text_cell = table.cell(0, 1)
            table.columns[0].width = img_col_width
            table.columns[1].width = text_col_width
        else:
            text_cell = table.cell(0, 0)
            img_cell = table.cell(0, 1)
            table.columns[0].width = text_col_width
            table.columns[1].width = img_col_width

        # Insert image in image cell
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(filepath, width=Inches(wrap_width_inches))

        # Add caption under image if provided
        if description:
            caption_para = img_cell.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption_para.add_run(description)
            cap_run.italic = True
            cap_run.font.size = Pt(8)
            cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Text cell gets a placeholder that will flow with surrounding content
        text_cell.paragraphs[0].text = ""

        # Set vertical alignment to top for both cells
        for cell in [img_cell, text_cell]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>')
            tcPr.append(vAlign)

    @classmethod
    def _insert_image_inline(cls, doc: Document, filepath: str, description: str):
        """Insert a small inline image (icon/logo size)."""
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(filepath, height=Inches(0.5))

    @classmethod
    def add_chapter_content(cls, doc: Document, title: str, content: str,
                            numbering: str, level: int = 1,
                            images: Optional[List[dict]] = None,
                            image_lookup: Optional[Dict[str, dict]] = None):
        """Add a chapter with content parsed from markdown.

        Args:
            image_lookup: Optional dict mapping image IDs to image info dicts
                          (with file_path, description keys). Used to resolve
                          [INSERT_IMAGE:id] markers in the generated content.
        """
        heading_text = f"{numbering} {title}" if numbering else title
        doc.add_heading(heading_text, level=min(level, 3))

        if content and content.strip():
            cls._parse_markdown_to_docx(
                doc, content,
                base_heading_level=min(level + 1, 3),
                image_lookup=image_lookup,
            )
        else:
            p = doc.add_paragraph("[Section à compléter]")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Add trailing images (legacy behavior for explicitly attached images)
        if images:
            for img_info in images:
                cls._insert_image_in_doc(
                    doc,
                    img_info.get("file_path", ""),
                    img_info.get("description", ""),
                )

    @classmethod
    def _parse_markdown_to_docx(cls, doc: Document, content: str,
                                base_heading_level: int = 2,
                                image_lookup: Optional[Dict[str, dict]] = None):
        """Parse markdown content and add properly formatted elements to the Word document.

        Handles: headings (##), bold/italic, bullet lists (- *), numbered lists (1.),
        tables (|...|), horizontal rules (---), [INSERT_IMAGE:id] markers,
        and regular paragraphs.

        Args:
            image_lookup: Optional dict mapping image IDs to image info dicts.
                          When a [INSERT_IMAGE:id] marker is encountered, the
                          corresponding original image is inserted into the document.
        """
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # [INSERT_IMAGE:id] or [INSERT_IMAGE:id:layout] marker
            img_match = _IMAGE_MARKER_RE.match(stripped)
            if img_match:
                img_id = img_match.group(1).strip()
                img_layout = (img_match.group(2) or "center").strip()
                if image_lookup and img_id in image_lookup:
                    img_info = image_lookup[img_id]
                    cls._insert_image_in_doc(
                        doc,
                        img_info.get("file_path", ""),
                        img_info.get("description", ""),
                        layout=img_layout,
                    )
                else:
                    logger.warning("Image marker [INSERT_IMAGE:%s] not found in lookup", img_id)
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
                # Add a thin horizontal line
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                i += 1
                continue

            # Markdown headings (## Title)
            heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if heading_match:
                hashes = heading_match.group(1)
                heading_text = heading_match.group(2).strip()
                md_level = len(hashes)
                # Map markdown level: ## -> base, ### -> base+1, etc.
                doc_level = min(base_heading_level + md_level - 2, 4)
                doc_level = max(1, doc_level)
                h = doc.add_heading(level=min(doc_level, 3))
                cls._add_inline_formatting(h, heading_text)
                i += 1
                continue

            # Table (lines starting with |)
            if stripped.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                cls._add_markdown_table(doc, table_lines)
                continue

            # Bullet list (- item or * item)
            if re.match(r'^[\-\*]\s+', stripped):
                while i < len(lines) and re.match(r'^\s*[\-\*]\s+', lines[i]):
                    item_text = re.sub(r'^\s*[\-\*]\s+', '', lines[i]).strip()
                    p = doc.add_paragraph(style='List Bullet')
                    cls._add_inline_formatting(p, item_text)
                    i += 1
                continue

            # Numbered list (1. item)
            if re.match(r'^\d+[\.\)]\s+', stripped):
                while i < len(lines) and re.match(r'^\s*\d+[\.\)]\s+', lines[i]):
                    item_text = re.sub(r'^\s*\d+[\.\)]\s+', '', lines[i]).strip()
                    p = doc.add_paragraph(style='List Number')
                    cls._add_inline_formatting(p, item_text)
                    i += 1
                continue

            # Regular paragraph - collect consecutive non-special lines
            para_lines = []
            while i < len(lines):
                l = lines[i].strip()
                if not l:
                    i += 1
                    break
                # Stop if next line is a special element
                if (re.match(r'^#{1,6}\s+', l) or
                    re.match(r'^[\-\*]\s+', l) or
                    re.match(r'^\d+[\.\)]\s+', l) or
                    l.startswith('|') or
                    re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', l) or
                    _IMAGE_MARKER_RE.match(l)):
                    break
                para_lines.append(lines[i].strip())
                i += 1

            if para_lines:
                para_text = ' '.join(para_lines)
                p = doc.add_paragraph()
                cls._add_inline_formatting(p, para_text)

    @classmethod
    async def generate_full_document(
        cls,
        project_name: str,
        client_name: str,
        rfp_reference: str,
        chapters: List[dict],
        company_name: str = "",
        image_lookup: Optional[Dict[str, dict]] = None,
    ) -> io.BytesIO:
        """Generate a complete Word document for the RFP response.

        Args:
            image_lookup: Optional dict mapping image IDs (str UUID) to image info
                          dicts with keys: file_path, description, image_type.
                          Used to resolve [INSERT_IMAGE:id] markers in chapter content.
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Create styles
        cls.create_styles(doc)

        # Cover page
        cls.add_cover_page(doc, project_name, client_name, rfp_reference, company_name)

        # Table of contents
        cls.add_table_of_contents(doc)

        # Add chapters
        def add_chapters_recursive(chapters_list: List[dict], level: int = 1, prefix: str = ""):
            for i, chapter in enumerate(chapters_list, 1):
                numbering = f"{prefix}{i}" if prefix else str(i)
                chapter["numbering"] = numbering

                cls.add_chapter_content(
                    doc,
                    title=chapter.get("title", ""),
                    content=chapter.get("content", ""),
                    numbering=numbering,
                    level=level,
                    images=chapter.get("images"),
                    image_lookup=image_lookup,
                )

                children = chapter.get("children", [])
                if children:
                    add_chapters_recursive(children, level + 1, f"{numbering}.")

                # Page break after major chapters
                if level == 1:
                    doc.add_page_break()

        # Separate main chapters from annexes
        main_chapters = [c for c in chapters if c.get("chapter_type") != "annexe"]
        annexes = [c for c in chapters if c.get("chapter_type") == "annexe"]

        add_chapters_recursive(main_chapters)

        # Add annexes section
        if annexes:
            doc.add_heading("ANNEXES", level=1)
            doc.add_paragraph()
            for i, annexe in enumerate(annexes, 1):
                cls.add_chapter_content(
                    doc,
                    title=annexe.get("title", ""),
                    content=annexe.get("content", ""),
                    numbering=f"A{i}",
                    level=2,
                    images=annexe.get("images"),
                    image_lookup=image_lookup,
                )

        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
